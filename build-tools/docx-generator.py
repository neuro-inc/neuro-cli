import argparse
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

PATH = Path.cwd()


def main(include: List[str], exclude: List[str], name: str) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(12)
    document.styles["Normal"].paragraph_format.line_spacing = 1.0
    document.styles["Normal"].paragraph_format.space_before = Pt(0)
    document.styles["Normal"].paragraph_format.space_after = Pt(10)
    document.styles["Heading 2"].font.name = "Arial"
    document.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    document.styles["Heading 2"].paragraph_format.space_after = Pt(10)

    def _check_exclude(path: Path) -> bool:
        path_str = str(path)
        for exclude_pattern in exclude:
            if exclude_pattern in path_str:
                return True
        return False

    for include_pattern in include:
        for path in PATH.glob(include_pattern):
            if _check_exclude(path):
                print(f"Skipping {path}")
                continue

            print(f"Adding {path}")
            document.add_heading(str(path.relative_to(PATH)), level=2)
            document.add_paragraph(path.read_text())

    for section in document.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    document.save(name)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("name", type=str)
    parser.add_argument("--include", action="append")
    parser.add_argument("--exclude", action="append")
    args = parser.parse_args()
    main(args.include or ["**/*"], args.exclude or [], args.name)
